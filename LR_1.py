from fuzzywuzzy import fuzz
from fuzzywuzzy import process

from docx import Document

PZ_1_p1_doc = Document()
PZ_1_p2_doc = Document()
PZ_1_p3_doc = Document()
PZ_1_p4_doc = Document()

p1 = PZ_1_p1_doc.add_paragraph('Программирование — это не просто работа, это настоящая страсть. Когда я пишу код, я чувствую себя творцом, способным создать что угодно из ничего. Любовь к программированию начинается с первой строчки, которая заставляет экран ожить, и продолжается в бессонных ночах ради исправления бага. Это искусство решать проблемы элегантно и логично, это бесконечный процесс обучения и совершенствования.')
p2 = PZ_1_p2_doc.add_paragraph('Программиравание — ето не проста работа, ето настоящая страсьть. Кагда я пишу код, я чуювствую себя творцом, спосоным создать что угодно из ничево. Любофь к праграмированию наченается с первой строчки, каторая заставляит экран ожить, и прадалжается в бессоных ночах ради исправления бага. Это исскуство решать праблемы элегантно и логичьно, это бесканечный працес абучения и совершеньствования.')
p3 = PZ_1_p3_doc.add_paragraph('Это страсть настоящая, программирование — не просто работа. Из ничего создать что угодно способным себя чувствую, когда код пишу я, творцом. К программированию любовь с первой строчки начинается, которая экран заставляет ожить. Бага исправления ради ночах бессонных продолжается и. Решать проблемы элегантно и логично — это искусство, это совершенствования процесс бесконечный и обучения.')
p4 = PZ_1_p4_doc.add_paragraph('Это искусство решать проблемы элегантно и логично, это бесконечный процесс обучения и совершенствования.')

p1_text = p1.text
p2_text = p2.text
p3_text = p3.text
p4_text = p4.text

PZ_1_p1_doc.save('PZ_1_p1_doc.docx')
PZ_1_p2_doc.save('PZ_1_p2_doc.docx')
PZ_1_p3_doc.save('PZ_1_p3_doc.docx')
PZ_1_p4_doc.save('PZ_1_p4_doc.docx')

def levenstein(str_1, str_2):
    n, m = len(str_1), len(str_2)
    if n > m:
        str_1, str_2 = str_2, str_1
        n, m = m, n

    current_row = range(n + 1)
    for i in range(1, m + 1):
        previous_row, current_row = current_row, [i] + [0] * n
        for j in range(1, n + 1):
            add, delete, change = previous_row[j] + 1, current_row[j - 1] + 1, previous_row[j - 1]
            if str_1[j - 1] != str_2[i - 1]:
                change += 1
            current_row[j] = min(add, delete, change)

    return current_row[n]

print(f"РЛ = {levenstein(p1_text, p2_text)}")

#Задачки по новым библиотекам

a1 = fuzz.ratio(p1_text, p1_text)
a2 = fuzz.ratio(p1_text, p2_text)
print(f"Обычное сравнение: {a1}, {a2}")

b1 = fuzz.partial_ratio(p4_text, p1_text)
b2 = fuzz.partial_ratio(p4_text, p2_text)
b3 = fuzz.partial_ratio(p4_text, p3_text)
print(f"Частичное сравнение: {b1}, {b2}, {b3}")

с1 = fuzz.token_sort_ratio(p1_text, p2_text)
с2 = fuzz.token_sort_ratio(p1_text, p3_text)
с3 = fuzz.token_sort_ratio(p1_text, p4_text)
print(f"Сравнение по токену: {с1}, {с2}, {с3}")

d1 = fuzz.WRatio(p1_text, p2_text)
d2 = fuzz.WRatio(p1_text, p3_text)
d3 = fuzz.WRatio(p1_text, p4_text)
print(f"Продвинутое обычное сравнение: {d1}, {d2}, {d3}")

city = ["Москва", "Санкт-Петербург", "Киров", "Киров!", "Киров!!!", "Кировск?"]
e1 = process.extract("Киров", city, limit=4)
e2 = process.extractOne("Краногрск", city)
print(f"Работа со списками: {e1}")
print(f"Работа со списками: {e2}")